import streamlit as st
import pandas as pd
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="AI와 함께하는 미래역량 특강 등록부")
st.header("특강 등록부 생성")

uploaded_file = st.file_uploader("설문 결과 CSV 파일을 업로드하세요.", type="csv")

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file)

    # 학번, 이름만 추출
    registration_df = df[['학번을 쓰시오', '이름을 쓰시오.']].copy()
    registration_df.columns = ['학번', '이름']

    # 학번 오름차순 정렬 (숫자/문자 혼합 처리)
    def 학번정렬키(x):
        try:
            return int(str(x).replace('-', ''))
        except:
            return str(x)
    registration_df = registration_df.sort_values(by='학번', key=lambda col: col.map(학번정렬키)).reset_index(drop=True)

    # 일련번호(구분) 추가
    registration_df.insert(0, '구분', range(1, len(registration_df)+1))
    # 서명, 비고 컬럼 추가
    registration_df['서명'] = ''
    registration_df['비고'] = ''
    # 컬럼 순서 재정렬
    registration_df = registration_df[['구분', '학번', '이름', '서명', '비고']]

    # --- 표 미리보기 (상위 10명만) ---
    st.subheader("등록부 미리보기 (상위 10명)")
    preview_df = registration_df.head(10)
    styled_preview = preview_df.style.set_properties(**{'text-align': 'center'}) \
        .set_table_styles([{'selector': 'th', 'props': [('text-align', 'center')]}])
    st.markdown(styled_preview.to_html(escape=False, index=False), unsafe_allow_html=True)

    # --- 엑셀(xlsx) 다운로드 ---
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
        registration_df.to_excel(writer, index=False, sheet_name='등록부')
        # 포맷 조정
        workbook  = writer.book
        worksheet = writer.sheets['등록부']
        # 컬럼 너비 조정
        worksheet.set_column('A:A', 8)
        worksheet.set_column('B:B', 16)
        worksheet.set_column('C:C', 16)
        worksheet.set_column('D:E', 18)
        # 헤더와 셀 가운데 정렬, 폰트 12pt
        header_format = workbook.add_format({'align': 'center', 'valign': 'vcenter', 'bold': True, 'font_size': 12})
        cell_format = workbook.add_format({'align': 'center', 'valign': 'vcenter', 'font_size': 12})
        for col_num, value in enumerate(registration_df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        worksheet.set_row(0, 24)
        for row_num in range(1, len(registration_df)+1):
            worksheet.set_row(row_num, 22)
            worksheet.write_row(row_num, 0, registration_df.iloc[row_num-1], cell_format)
    excel_buffer.seek(0)
    st.download_button(
        label="엑셀(xlsx)로 등록부 다운로드",
        data=excel_buffer,
        file_name="특강등록부.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    # --- 워드(docx) 다운로드 ---
    doc_buffer = io.BytesIO()
    doc = Document()
    doc.add_heading('AI와 함께하는 미래역량 특강 등록부', 0)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    columns = ['구분', '학번', '이름', '서명', '비고']
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(12)
    # 데이터
    for idx, row in registration_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0]
                run.font.size = Pt(12)
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    st.download_button(
        label="워드(docx)로 등록부 다운로드",
        data=doc_buffer,
        file_name="특강등록부.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.info("엑셀/워드 파일은 편집·인쇄가 모두 가능합니다. 표 서식(폰트, 셀 크기, 정렬)도 적용되어 있습니다.")
else:
    st.info("CSV 파일을 업로드하면 등록부 미리보기와 편집 가능한 엑셀/워드 다운로드가 가능합니다.")
